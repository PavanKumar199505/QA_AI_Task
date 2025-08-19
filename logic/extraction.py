import logging
from pathlib import Path
import pytesseract
import fitz  # PyMuPDF
from PIL import Image
import whisper
import subprocess
import tempfile
import os
from docx import Document
import pdfplumber
from pdf2image import convert_from_path
import zipfile
from lxml import etree

logger = logging.getLogger(__name__)


def extract_text_from_file(file_path: Path) -> str:
    """
    Extract text from various file types (image, PDF, audio, video, text, Word).

    Args:
        file_path: Path to the input file.

    Returns:
        Extracted text as a string.

    Raises:
        Exception: If extraction fails or file type is unsupported.
    """
    file_path = Path(file_path)
    extension = file_path.suffix.lower()

    logger.info(f"Attempting to extract text from {file_path} (type: {extension})")

    try:
        if extension in [".png", ".jpg", ".jpeg"]:
            return extract_text_from_image(file_path)
        elif extension == ".pdf":
            return extract_text_from_pdf(file_path)
        elif extension in [".mp3", ".wav"]:
            return extract_text_from_audio(file_path)
        elif extension == ".mp4":
            return extract_text_from_video(file_path)
        elif extension == ".txt":
            return extract_text_from_text(file_path)
        elif extension == ".docx":
            return extract_text_from_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")
    except Exception as e:
        logger.error(f"Error extracting text from {file_path}: {str(e)}", exc_info=True)
        raise


def extract_text_from_image(file_path: Path) -> str:
    """Extract text from an image using Tesseract OCR."""
    try:
        logger.info(f"Extracting text from image: {file_path}")
        image = Image.open(file_path)
        text = pytesseract.image_to_string(image, lang='eng', config='--psm 6')
        result = text.strip()
        logger.info(f"Successfully extracted {len(result)} characters from image")
        return result
    except Exception as e:
        logger.error(f"OCR failed for {file_path}: {str(e)}")
        raise Exception(f"OCR failed: {str(e)}. Ensure Tesseract is installed and accessible.")


def extract_text_from_pdf(file_path: Path) -> str:
    """Extract text from a PDF using multiple methods with fallbacks."""
    try:
        logger.info(f"Extracting text from PDF: {file_path}")
        
        # Method 1: Try pdfplumber first (better for text-based PDFs)
        try:
            with pdfplumber.open(file_path) as pdf:
                text_parts = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text and page_text.strip():
                        text_parts.append(page_text.strip())
                
                if text_parts:
                    result = "\n".join(text_parts)
                    logger.info(f"Successfully extracted {len(result)} characters using pdfplumber")
                    return result
        except Exception as e:
            logger.warning(f"pdfplumber failed: {str(e)}, trying PyMuPDF...")

        # Method 2: Try PyMuPDF
        try:
            doc = fitz.open(file_path)
            text_parts = []
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                page_text = page.get_text().strip()
                if page_text:
                    text_parts.append(page_text)
                else:
                    # Try OCR for image-based pages
                    pix = page.get_pixmap()
                    img = Image.frombytes("RGB", [pix.width, pix.height], pix.rgb)
                    page_text = pytesseract.image_to_string(img, lang='eng', config='--psm 6')
                    if page_text.strip():
                        text_parts.append(page_text.strip())
            doc.close()
            
            if text_parts:
                result = "\n".join(text_parts)
                logger.info(f"Successfully extracted {len(result)} characters using PyMuPDF")
                return result
        except Exception as e:
            logger.warning(f"PyMuPDF failed: {str(e)}, trying OCR...")

        # Method 3: OCR fallback
        try:
            images = convert_from_path(file_path, dpi=300)
            text_parts = []
            for image in images:
                text = pytesseract.image_to_string(image, lang='eng', config='--psm 6')
                if text.strip():
                    text_parts.append(text.strip())
            
            if text_parts:
                result = "\n".join(text_parts)
                logger.info(f"Successfully extracted {len(result)} characters using OCR")
                return result
        except Exception as e:
            logger.error(f"OCR fallback failed: {str(e)}")

        raise Exception("All PDF extraction methods failed. The PDF might be encrypted, corrupted, or contain no extractable text.")
        
    except Exception as e:
        logger.error(f"PDF extraction failed for {file_path}: {str(e)}")
        raise Exception(f"PDF extraction failed: {str(e)}")


def extract_text_from_audio(file_path: Path) -> str:
    """Transcribe audio to text using Whisper."""
    try:
        logger.info(f"Transcribing audio: {file_path}")
        model = whisper.load_model("base")
        result = model.transcribe(str(file_path))
        text = result["text"].strip()
        logger.info(f"Successfully transcribed {len(text)} characters from audio")
        return text
    except Exception as e:
        logger.error(f"Audio transcription failed for {file_path}: {str(e)}")
        raise Exception(f"Audio transcription failed: {str(e)}. Ensure Whisper and FFmpeg are installed.")


def extract_text_from_video(file_path: Path) -> str:
    """Extract audio from video in chunks and transcribe to text."""
    try:
        logger.info(f"Extracting audio from video: {file_path}")
        
        # Get video duration using ffprobe
        result = subprocess.run(
            [
                "ffprobe", "-v", "error", "-show_entries", "format=duration",
                "-of", "default=noprint_wrappers=1:nokey=1", str(file_path)
            ],
            capture_output=True, text=True, check=True
        )
        duration = float(result.stdout)
        
        logger.info(f"Video duration: {duration:.2f} seconds. Processing in 60-second chunks.")
        
        model = whisper.load_model("base")
        full_text = []
        
        # Process in 60-second chunks
        for start_time in range(0, int(duration), 60):
            chunk_duration = min(60, duration - start_time)
            logger.info(f"Processing chunk: {start_time}s - {start_time + chunk_duration}s")
            
            with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as temp_audio:
                temp_audio_path = temp_audio.name
            
            try:
                # Extract audio chunk using FFmpeg
                subprocess.run(
                    [
                        "ffmpeg", "-i", str(file_path), "-ss", str(start_time),
                        "-t", str(chunk_duration), "-vn", "-acodec", "pcm_s16le",
                        "-ar", "16000", "-ac", "1", temp_audio_path
                    ],
                    check=True, capture_output=True, text=True
                )
                
                # Transcribe the chunk
                transcription_result = model.transcribe(temp_audio_path)
                chunk_text = transcription_result["text"].strip()
                if chunk_text:
                    full_text.append(chunk_text)
            
            finally:
                # Clean up the temporary file
                try:
                    os.remove(temp_audio_path)
                except OSError:
                    pass
        
        final_text = " ".join(full_text)
        logger.info(f"Successfully transcribed {len(final_text)} characters from video.")
        return final_text

    except subprocess.CalledProcessError as e:
        logger.error(f"FFmpeg or ffprobe failed: {e.stderr}")
        raise Exception("Video processing failed. Ensure FFmpeg is installed and accessible.")
    except Exception as e:
        logger.error(f"Video transcription failed for {file_path}: {str(e)}")
        raise Exception(f"Video transcription failed: {str(e)}")


def extract_text_from_text(file_path: Path) -> str:
    """Read text from a plain text file."""
    try:
        logger.info(f"Reading text file: {file_path}")
        with open(file_path, "r", encoding="utf-8") as f:
            text = f.read().strip()
        logger.info(f"Successfully read {len(text)} characters from text file")
        return text
    except UnicodeDecodeError:
        # Try with different encoding
        try:
            with open(file_path, "r", encoding="latin-1") as f:
                text = f.read().strip()
            logger.info(f"Successfully read {len(text)} characters using latin-1 encoding")
            return text
        except Exception as e:
            logger.error(f"Text file reading failed with latin-1 encoding: {str(e)}")
            raise Exception(f"Text file reading failed: Unable to decode file with UTF-8 or latin-1 encoding.")
    except Exception as e:
        logger.error(f"Text file reading failed for {file_path}: {str(e)}")
        raise Exception(f"Text file reading failed: {str(e)}")


def extract_text_from_docx(file_path: Path) -> str:
    """Extract text from a Word document with enhanced error handling."""
    try:
        logger.info(f"Extracting text from DOCX: {file_path}")
        
        # Validate if it's a valid DOCX (ZIP archive)
        try:
            with zipfile.ZipFile(file_path) as z:
                if 'word/document.xml' not in z.namelist():
                    raise ValueError("Invalid DOCX file: missing document.xml")
        except zipfile.BadZipFile:
            logger.warning(f"File {file_path} is not a valid ZIP archive, trying raw XML extraction...")
            return extract_text_from_raw_xml(file_path)
        
        # Standard DOCX extraction
        doc = Document(file_path)
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text.strip())
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        
        # Extract from headers and footers
        for section in doc.sections:
            for header in section.header.paragraphs:
                if header.text.strip():
                    text_parts.append(header.text.strip())
            for footer in section.footer.paragraphs:
                if footer.text.strip():
                    text_parts.append(footer.text.strip())
        
        result = "\n".join(text_parts)
        logger.info(f"Successfully extracted {len(result)} characters from DOCX")
        return result
        
    except Exception as e:
        logger.error(f"DOCX extraction failed for {file_path}: {str(e)}")
        raise Exception(f"Word document extraction failed: {str(e)}")


def extract_text_from_raw_xml(file_path: Path) -> str:
    """Extract text from raw Word document.xml with external entity resolution disabled."""
    try:
        logger.info(f"Extracting text from raw XML: {file_path}")
        parser = etree.XMLParser(resolve_entities=False, load_dtd=False, no_network=True)
        with open(file_path, 'rb') as f:
            tree = etree.parse(f, parser)

        # Define namespaces
        namespaces = {
            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
            'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
            'wps': 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
        }

        text_parts = []
        # Extract text from paragraphs (<w:t> elements)
        for elem in tree.xpath('//w:t', namespaces=namespaces):
            if elem.text and elem.text.strip():
                text_parts.append(elem.text.strip())

        # Extract text from shapes or text boxes (wps:t elements)
        for elem in tree.xpath('//wps:t', namespaces=namespaces):
            if elem.text and elem.text.strip():
                text_parts.append(elem.text.strip())

        result = "\n".join(text_parts)
        if not result:
            raise ValueError("No readable text found in XML document")
        
        logger.info(f"Successfully extracted {len(result)} characters from raw XML")
        return result

    except etree.ParseError as e:
        logger.error(f"XML parsing failed: {str(e)}")
        raise Exception(f"Failed to parse XML: {str(e)}")
    except Exception as e:
        logger.error(f"Raw XML extraction failed: {str(e)}")
        raise Exception(f"Failed to extract text from XML: {str(e)}")